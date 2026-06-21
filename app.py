import os
import io
import json
import tempfile
import requests as http_requests
from flask import Flask, request, jsonify, send_file, render_template
from flask_cors import CORS
from dotenv import load_dotenv
from pypdf import PdfReader
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from nlp_engine import compute_ats_score

load_dotenv()

app = Flask(__name__)
CORS(app)

# ── LOCAL OLLAMA CONFIG ───────────────────────────────────────────────────────
OLLAMA_URL   = os.getenv("OLLAMA_URL", "http://localhost:11434")
OLLAMA_MODEL = os.getenv("OLLAMA_MODEL", "qwen2.5:7b")

ALLOWED_EXTENSIONS = {"pdf", "docx"}
MAX_FILE_SIZE = 5 * 1024 * 1024  # 5MB


# ── UTILS ─────────────────────────────────────────────────────────────────────

def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


def extract_text_from_pdf(file_bytes):
    reader = PdfReader(io.BytesIO(file_bytes))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def extract_text_from_docx(file_bytes):
    doc = Document(io.BytesIO(file_bytes))
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def extract_resume_text(file_bytes, filename):
    ext = filename.rsplit(".", 1)[1].lower()
    if ext == "pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext == "docx":
        return extract_text_from_docx(file_bytes)
    return ""


# ── LLM ───────────────────────────────────────────────────────────────────────

SYSTEM_PROMPT = """You are an expert ATS resume writer. Given a candidate's details, their existing resume, and a job description, generate a perfectly tailored, ATS-optimized resume.

Return ONLY a valid JSON object with this exact structure (no markdown, no extra text):
{
  "name": "",
  "email": "",
  "phone": "",
  "location": "",
  "linkedin": "",
  "github": "",
  "summary": "",
  "skills": ["skill1", "skill2"],
  "experience": [
    {
      "title": "",
      "company": "",
      "duration": "",
      "bullets": ["bullet1", "bullet2", "bullet3"]
    }
  ],
  "education": [
    {
      "degree": "",
      "institution": "",
      "year": "",
      "gpa": ""
    }
  ],
  "projects": [
    {
      "name": "",
      "tech": "",
      "bullets": ["bullet1", "bullet2"]
    }
  ],
  "certifications": ["cert1", "cert2"],
  "matched_skills": ["Python", "TensorFlow"],
  "missing_skills": ["Docker", "Kubernetes"]
}

Rules:
- Mirror keywords from the job description naturally throughout
- Use strong action verbs for all bullet points
- Quantify achievements wherever possible
- Keep summary to 2-3 sentences, ATS-keyword rich
- Skills must directly match job description requirements
- Do NOT invent experience that isn't in the original resume
- Return ONLY the JSON, nothing else
"""


def generate_resume_json(details, job_description, existing_resume_text):
    user_prompt = f"""CANDIDATE DETAILS:
Name: {details.get('name', '')}
Email: {details.get('email', '')}
Phone: {details.get('phone', '')}
Location: {details.get('location', '')}
LinkedIn: {details.get('linkedin', '')}
GitHub: {details.get('github', '')}
Target Role: {details.get('targetRole', '')}

JOB DESCRIPTION:
{job_description}

EXISTING RESUME:
{existing_resume_text}"""

    payload = {
        "model": OLLAMA_MODEL,
        "messages": [
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user",   "content": user_prompt},
        ],
        "stream": False,
        "options": {
            "temperature": 0.3,   # lower = more deterministic JSON output
            "num_predict": 4096,  # enough tokens for a full resume JSON
        },
    }

    resp = http_requests.post(
        f"{OLLAMA_URL}/api/chat",
        json=payload,
        timeout=300,  # local inference can be slow — 5 min ceiling
    )
    resp.raise_for_status()

    raw = resp.json()["message"]["content"].strip()

    # Strip markdown fences if the model wraps output in ```json ... ```
    if raw.startswith("```"):
        raw = raw.split("```")[1]
        if raw.startswith("json"):
            raw = raw[4:]
    return json.loads(raw.strip())


# ── DOCX GENERATOR ────────────────────────────────────────────────────────────

def build_docx(data):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = section.bottom_margin = 914400 // 2   # 0.5 inch
        section.left_margin = section.right_margin = 914400 // 1   # 1 inch

    def heading1(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = 16 * 12700  # 16pt
        p.paragraph_format.space_after = 60960

    def section_title(text):
        p = doc.add_paragraph()
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = 11 * 12700
        p.paragraph_format.space_before = 120000
        p.paragraph_format.space_after = 60960
        doc.add_paragraph("─" * 60)

    def body(text, bold=False):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.font.size = 10 * 12700
        p.paragraph_format.space_after = 30000

    def bullet(text):
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(text)
        run.font.size = 10 * 12700
        p.paragraph_format.space_after = 20000

    # Header
    heading1(data.get("name", ""))
    contact = " | ".join(filter(None, [
        data.get("email"), data.get("phone"), data.get("location"),
        data.get("linkedin"), data.get("github")
    ]))
    body(contact)

    # Summary
    if data.get("summary"):
        section_title("Professional Summary")
        body(data["summary"])

    # Skills
    if data.get("skills"):
        section_title("Skills")
        body(", ".join(data["skills"]))

    # Experience
    if data.get("experience"):
        section_title("Experience")
        for exp in data["experience"]:
            body(f"{exp.get('title','')} — {exp.get('company','')}  |  {exp.get('duration','')}", bold=True)
            for b in exp.get("bullets", []):
                bullet(b)

    # Projects
    if data.get("projects"):
        section_title("Projects")
        for proj in data["projects"]:
            body(f"{proj.get('name','')}  |  {proj.get('tech','')}", bold=True)
            for b in proj.get("bullets", []):
                bullet(b)

    # Education
    if data.get("education"):
        section_title("Education")
        for edu in data["education"]:
            gpa = f"  |  GPA: {edu['gpa']}" if edu.get("gpa") else ""
            body(f"{edu.get('degree','')} — {edu.get('institution','')}  |  {edu.get('year','')}{gpa}", bold=True)

    # Certifications
    if data.get("certifications"):
        section_title("Certifications")
        for cert in data["certifications"]:
            bullet(cert)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ── PDF GENERATOR ─────────────────────────────────────────────────────────────

def build_pdf(data):
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        rightMargin=18*mm, leftMargin=18*mm,
        topMargin=16*mm, bottomMargin=16*mm
    )

    styles = getSampleStyleSheet()
    ACCENT = colors.HexColor("#2d3a8c")
    GRAY   = colors.HexColor("#555555")

    s_name    = ParagraphStyle("name",    fontSize=20, fontName="Helvetica-Bold", textColor=ACCENT, spaceAfter=2)
    s_contact = ParagraphStyle("contact", fontSize=9,  fontName="Helvetica",      textColor=GRAY,   spaceAfter=8)
    s_sec     = ParagraphStyle("sec",     fontSize=11, fontName="Helvetica-Bold", textColor=ACCENT, spaceBefore=10, spaceAfter=2)
    s_body    = ParagraphStyle("body",    fontSize=9.5, fontName="Helvetica",     leading=14,       spaceAfter=3)
    s_bold    = ParagraphStyle("bold",    fontSize=9.5, fontName="Helvetica-Bold",leading=14,       spaceAfter=2)
    s_bullet  = ParagraphStyle("bul",     fontSize=9.5, fontName="Helvetica",     leading=14,       leftIndent=12, bulletIndent=2, spaceAfter=2)

    def hr():
        return HRFlowable(width="100%", thickness=0.5, color=ACCENT, spaceAfter=4)

    def sec(title):
        return [Paragraph(title.upper(), s_sec), hr()]

    story = []

    # Header
    story.append(Paragraph(data.get("name", ""), s_name))
    contact = " · ".join(filter(None, [
        data.get("email"), data.get("phone"), data.get("location"),
        data.get("linkedin"), data.get("github")
    ]))
    story.append(Paragraph(contact, s_contact))

    # Summary
    if data.get("summary"):
        story += sec("Professional Summary")
        story.append(Paragraph(data["summary"], s_body))
        story.append(Spacer(1, 4))

    # Skills
    if data.get("skills"):
        story += sec("Skills")
        story.append(Paragraph(", ".join(data["skills"]), s_body))
        story.append(Spacer(1, 4))

    # Experience
    if data.get("experience"):
        story += sec("Experience")
        for exp in data["experience"]:
            story.append(Paragraph(f"{exp.get('title','')} — {exp.get('company','')} | {exp.get('duration','')}", s_bold))
            for b in exp.get("bullets", []):
                story.append(Paragraph(f"• {b}", s_bullet))
        story.append(Spacer(1, 4))

    # Projects
    if data.get("projects"):
        story += sec("Projects")
        for proj in data["projects"]:
            story.append(Paragraph(f"{proj.get('name','')} | {proj.get('tech','')}", s_bold))
            for b in proj.get("bullets", []):
                story.append(Paragraph(f"• {b}", s_bullet))
        story.append(Spacer(1, 4))

    # Education
    if data.get("education"):
        story += sec("Education")
        for edu in data["education"]:
            gpa = f" | GPA: {edu['gpa']}" if edu.get("gpa") else ""
            story.append(Paragraph(f"{edu.get('degree','')} — {edu.get('institution','')} | {edu.get('year','')}{gpa}", s_bold))
        story.append(Spacer(1, 4))

    # Certifications
    if data.get("certifications"):
        story += sec("Certifications")
        for cert in data["certifications"]:
            story.append(Paragraph(f"• {cert}", s_bullet))

    doc.build(story)
    buf.seek(0)
    return buf


# ── ROUTES ────────────────────────────────────────────────────────────────────

@app.route("/api/health", methods=["GET"])
def health():
    return jsonify({"status": "ok"})


@app.route("/api/generate", methods=["POST"])
def generate():
    # Validate inputs
    if "resume" not in request.files:
        return jsonify({"error": "No resume file uploaded"}), 400

    file = request.files["resume"]
    if not file or not allowed_file(file.filename):
        return jsonify({"error": "Invalid file type. Use PDF or DOCX"}), 400

    file_bytes = file.read()
    if len(file_bytes) > MAX_FILE_SIZE:
        return jsonify({"error": "File too large. Max 5MB"}), 400

    job_description = request.form.get("jobDescription", "").strip()
    if not job_description:
        return jsonify({"error": "Job description is required"}), 400

    details = {
        "name":       request.form.get("name", ""),
        "email":      request.form.get("email", ""),
        "phone":      request.form.get("phone", ""),
        "location":   request.form.get("location", ""),
        "linkedin":   request.form.get("linkedin", ""),
        "github":     request.form.get("github", ""),
        "targetRole": request.form.get("targetRole", ""),
    }

    try:
        existing_text = extract_resume_text(file_bytes, file.filename)
    except Exception as e:
        return jsonify({"error": f"Failed to parse resume: {str(e)}"}), 422

    try:
        resume_data = generate_resume_json(details, job_description, existing_text)
    except json.JSONDecodeError:
        return jsonify({"error": "LLM returned invalid JSON. Try again."}), 500
    except Exception as e:
        return jsonify({"error": f"LLM error: {str(e)}"}), 500

    try:
        pdf_buf  = build_pdf(resume_data)
        docx_buf = build_docx(resume_data)
    except Exception as e:
        return jsonify({"error": f"Document generation failed: {str(e)}"}), 500

    # Save to temp files and return paths via JSON
    # (Frontend will hit /api/download/pdf and /api/download/docx)
    tmp_dir = tempfile.gettempdir()
    pdf_path  = os.path.join(tmp_dir, "resume_output.pdf")
    docx_path = os.path.join(tmp_dir, "resume_output.docx")

    with open(pdf_path, "wb") as f:
        f.write(pdf_buf.getvalue())
    with open(docx_path, "wb") as f:
        f.write(docx_buf.getvalue())

    # ── Phase 2: ATS Scoring ───────────────────────────────────────────────
    try:
        # Build full resume text for NLP comparison
        resume_full_text = " ".join([
            resume_data.get("summary", ""),
            " ".join(resume_data.get("skills", [])),
            " ".join(
                b for exp in resume_data.get("experience", [])
                for b in exp.get("bullets", [])
            ),
            " ".join(
                b for proj in resume_data.get("projects", [])
                for b in proj.get("bullets", [])
            ),
        ])
        ats = compute_ats_score(
            jd_text=job_description,
            resume_text=resume_full_text,
            matched_skills=resume_data.get("matched_skills", []),
            missing_skills=resume_data.get("missing_skills", []),
        )
    except Exception as e:
        # NLP scoring is non-critical — degrade gracefully
        ats = {
            "ats_score": 0, "keyword_coverage": 0,
            "skills_match": 0, "experience_relevance": 0,
            "keyword_hits": [], "keyword_misses": [],
        }

    return jsonify({
        "status":               "success",
        "name":                 resume_data.get("name", "resume"),
        "matched_skills":       resume_data.get("matched_skills", []),
        "missing_skills":       resume_data.get("missing_skills", []),
        "ats_score":            ats["ats_score"],
        "keyword_coverage":     ats["keyword_coverage"],
        "skills_match":         ats["skills_match"],
        "experience_relevance": ats["experience_relevance"],
        "keyword_hits":         ats["keyword_hits"],
        "keyword_misses":       ats["keyword_misses"],
        "interview_questions":  [],
        "cover_letter_ready":   False,
    })

@app.route("/api/download/pdf", methods=["GET"])
def download_pdf():
    path = os.path.join(tempfile.gettempdir(), "resume_output.pdf")
    if not os.path.exists(path):
        return jsonify({"error": "File not found. Generate resume first."}), 404
    return send_file(path, mimetype="application/pdf",
                     as_attachment=True, download_name="resume.pdf")


@app.route("/api/download/docx", methods=["GET"])
def download_docx():
    path = os.path.join(tempfile.gettempdir(), "resume_output.docx")
    if not os.path.exists(path):
        return jsonify({"error": "File not found. Generate resume first."}), 404
    return send_file(path,
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                     as_attachment=True, download_name="resume.docx")


# ── MAIN ──────────────────────────────────────────────────────────────────────


@app.route("/")
def home():
    return render_template("index.html")

if __name__ == "__main__":
    app.run(debug=True, port=5000)